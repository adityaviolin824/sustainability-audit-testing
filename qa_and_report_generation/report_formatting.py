import json
import sys
from pathlib import Path
from typing import Dict, List, Any

from docx import Document
from openpyxl import Workbook
from utils.logger import logging
from utils.exception import CustomException

logger = logging.getLogger(__name__)

# --------------------------------------------------
# Helpers
# --------------------------------------------------

def infer_pillar(metric_id: str) -> str:
    """Categorizes metrics into ESG pillars based on their ID prefix."""
    if metric_id.startswith("ENV"):
        return "Environmental"
    if metric_id.startswith(("SRC", "WRK", "SOC", "EWB", "HR", "CSR", "STK", "CST")):
        return "Social"
    if metric_id.startswith("GOV"):
        return "Governance"
    return "Other"

def safe_text(value: Any) -> str:
    """Returns a fallback string if the value is empty or None."""
    return value if value not in (None, "", []) else "Not disclosed in the report."

def load_questions_mapping(path: Path) -> Dict[str, str]:
    """Loads JSONL questions into a lookup dictionary {metric_id: question_text}."""
    try:
        mapping = {}
        with open(path, "r", encoding="utf-8") as f:
            for line in f:
                if not line.strip(): continue
                obj = json.loads(line)
                mapping[obj["id"]] = obj["question"]
        return mapping
    except Exception as e:
        raise CustomException(e, sys)

# --------------------------------------------------
# Document Generation Modules
# --------------------------------------------------

def generate_docx_report(records: List[Dict], questions_map: Dict[str, str], output_path: Path):
    """Generates a formatted Word document for the ESG Audit."""
    try:
        doc = Document()
        doc.add_heading("ESG Audit – Extracted Disclosures", level=0)

        current_pillar = None

        for rec in records:
            metric_id = rec["metric_id"]
            pillar = infer_pillar(metric_id)

            # Add Section Header if Pillar changes
            if pillar != current_pillar:
                doc.add_heading(pillar, level=1)
                current_pillar = pillar

            doc.add_heading(f"Metric ID: {metric_id}", level=2)

            # Question Section
            p_q = doc.add_paragraph()
            p_q.add_run("Question:").bold = True
            doc.add_paragraph(questions_map.get(metric_id, "Question not found in question bank."))

            # Summary Section
            p_s = doc.add_paragraph()
            p_s.add_run("Summary:").bold = True
            doc.add_paragraph(safe_text(rec.get("summary")))

            # Key Facts Section
            p_f = doc.add_paragraph()
            p_f.add_run("Key Facts:").bold = True
            key_facts = rec.get("key_facts", [])
            if key_facts:
                if isinstance(key_facts, list):
                    for fact in key_facts:
                        doc.add_paragraph(str(fact), style="List Bullet")
                else:
                    doc.add_paragraph(str(key_facts))
            else:
                doc.add_paragraph("Not disclosed in the report.")

            # Page Reference
            p_p = doc.add_paragraph()
            p_p.add_run("Page Reference:").bold = True
            doc.add_paragraph(safe_text(rec.get("page")))

            doc.add_paragraph("-" * 30)

        doc.save(output_path)
        logger.info(f"Word report generated: {output_path}")
    except Exception as e:
        raise CustomException(e, sys)


def generate_xlsx_report(records: List[Dict], output_path: Path):
    """Generates a structured Excel file for ESG data analysis."""
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "ESG Disclosures"

        # Define Headers
        ws.append(["Metric ID", "Pillar", "Summary", "Key Facts", "Page"])

        for rec in records:
            key_facts = rec.get("key_facts", [])
            
            # Type check to prevent .join() from iterating over a single string
            if isinstance(key_facts, list):
                key_facts_flat = "; ".join(map(str, key_facts)) if key_facts else "Not disclosed in the report."
            else:
                key_facts_flat = str(key_facts) if key_facts else "Not disclosed in the report."

            ws.append([
                rec.get("metric_id"),
                infer_pillar(rec.get("metric_id", "")),
                safe_text(rec.get("summary")),
                key_facts_flat,
                safe_text(rec.get("page")),
            ])

        wb.save(output_path)
        logger.info(f"Excel report generated: {output_path}")
    except Exception as e:
        raise CustomException(e, sys)

# --------------------------------------------------
# Pipeline Orchestrator
# --------------------------------------------------

def run_reporting_pipeline(
    consolidated_json_path: Path,
    questions_jsonl_path: Path,
    output_docx_path: Path,
    output_xlsx_path: Path
):
    """Executes the final reporting logic from data files."""
    try:
        if not consolidated_json_path.exists():
            raise FileNotFoundError(f"Consolidated data not found: {consolidated_json_path}")

        # Load Data
        logger.info("Loading consolidated records and question mapping...")
        records = json.loads(consolidated_json_path.read_text(encoding="utf-8"))
        questions_map = load_questions_mapping(questions_jsonl_path)

        # Generate Artifacts
        generate_docx_report(records, questions_map, output_docx_path)
        generate_xlsx_report(records, output_xlsx_path)

        logger.info(">>> Reporting pipeline execution successful <<<")

    except Exception as e:
        raise CustomException(e, sys)
